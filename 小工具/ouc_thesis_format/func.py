# -*- coding: utf-8 -*-
"""
Created on Sun Mar 29 23:05:25 2020

@author: deng
"""
import win32com.client
import docx
from docx.oxml.ns import qn  # 设置中文字体
from docx.shared import Pt   # 字体大小


# 论文信息
info = {
    '姓名': '张三',
    '论文题目': '这是论文题目',
    '英文题目': '这是英文题目',
    '指导教师': '李四'
}


def set_char_name(run, name, size):
    run.font.size = Pt(size)
    run.font.name = name
    # 设置字体必须要下面2步
    s = run._element
    s.rPr.rFonts.set(qn('w:eastAsia'), name)


doc = docx.Document('封皮.docx')  # 打开文件


# =============================================================================
# # 获取文本框文本
# #遍历文档每个子节点
# for child in doc.element.body.iter():
# #只处理textbox
#     if child.tag.endswith('textbox'):
#     #if child.tag.endswith('main}std'):
#         print(f'\n{"="*20}',end='')
#         #遍历文本框中每个子节点
#         for c in child.iter():
#             c_tag=c.tag
#             #遇到段落,换行
#             if c_tag.endswith('main}pPr'):
#                 print()
#             #遇到段内run,提取并输出文本
#             elif c_tag.endswith('main}r'):
#                 print(c.text,end='')
# =============================================================================


# 谨以此论文献给我的导师、亲人和父母
doc.paragraphs[2].text = f'                                 ------{info["姓名"]}'
set_char_name(doc.paragraphs[2].runs[0], '黑体', 18)  # 小二

# 论文题目
doc.paragraphs[4].text = info['论文题目']
set_char_name(doc.paragraphs[4].runs[0], '黑体', 18)  # 小二

# 添加正文
# doc.add_page_break()   # 添加换页符
sub_doc = docx.Document('正文.docx')
# =============================================================================
# for i in range(len(doc.paragraphs)):
#     if doc.paragraphs[i].text == '签字日期：    年   月   日                签字日期：    年   月   日': print(i)
# =============================================================================
n = 35
doc.paragraphs[n].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)  # 指定位置添加换页符
for element in sub_doc.element.body:
    doc.paragraphs[n]._p.addnext(element)
    n += 1
    # doc.element.body.append(element)

# 更改页眉，引言在正文第7页
i = 11
doc.sections[i-1].header.is_linked_to_previous = False
doc.sections[i].header.paragraphs[0].style = doc.sections[1].header.paragraphs[0].style
doc.sections[i].header.paragraphs[0].text = f"\n{info['论文题目']}"
doc.sections[1].header.is_linked_to_previous = True


# 保存结果
doc.save('output.docx')

# 更新目录
word = win32com.client.DispatchEx("Word.Application")
doc = word.Documents.Open('output.docx')
doc.TablesOfContents(1).Update()
doc.Close(SaveChanges=True)
word.Quit()
